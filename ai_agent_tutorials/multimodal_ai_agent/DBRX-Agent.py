from turtle import setx
import streamlit as st
from agno.agent import Agent
from agno.models.google import Gemini
from agno.models.together import Together
from agno.models.openai.like import OpenAILike
from phi.tools.duckduckgo import DuckDuckGo
from dotenv import load_dotenv
import time
import os
from pathlib import Path
import tempfile
from docx import Document
import markdown

load_dotenv()
import google.generativeai as genai
from google.generativeai import upload_file, get_file
genai.configure(api_key=os.environ['GOOGLE_API_KEY'])

# New together.ai API integration (placeholder)
# TOGETHER_API_KEY = os.getenv("TOGETHER_API_KEY")

agent = Agent(
        # model=Gemini(id="gemini-2.0-flash-thinking-exp-1219"), 
        # model=Gemini(id="gemini-2.0-flash-thinking-exp"),
        # model=Gemini(id="gemini-2.0-flash"),
        # model=Gemini(id="gemini-2.0-flash-001"),
        model = OpenAILike(
            # id="deepseek-ai/DeepSeek-R1",
            id="databricks/dbrx-instruct",
            api_key = os.getenv("TOGETHER_API_KEY"),
            base_url="https://api.together.xyz/v1",
            ),
        # model=Gemini(id="gemini-2.0-pro-exp-02-05"),
            markdown=True,
           )
       
def generate_response(task_input):
    # Call the agent with the dynamic task and image path
    response = agent.run(task_input)
    return response

def create_word_document(content):
    doc = Document()
    doc.add_heading('AI Response', level=1)
    doc.add_paragraph(content)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp.seek(0)
        return tmp.read()

def create_html_content(content):
    return markdown.markdown(content)

# Custom CSS
st.markdown(
    """
    <style>
    .title {
        color: crimson;
        margin-top: -75px; /* Adjust this value to reduce spacing */
        border-bottom: 2px solid darkslategray;
        padding-bottom: 5px;
        font-size: 35px; /* Adjust this value to change the font size */
    }
    .stButton button {
        background-color: indianred;
        color: white;
    }
    .ai-response-heading {
        color: darkslategray;
        font-size: 23px; /* Adjust this value to change the font size */
    }
    .stDownloadButton button {
        background-color: indianred;
        color: white;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Main Streamlit app
def main():
    st.markdown('<h1 class="title">Databricks Agent</h1>', unsafe_allow_html=True)
    task_input = st.text_area("Enter your task/question for the AI Agent:")

    if st.button("Develop Response") and task_input:
        with st.spinner("AI is thinking... 🤖"):
            try:
                response = generate_response(task_input)
                st.session_state['response_content'] = response.content
            except Exception as e:
                st.error(f"An error occurred during analysis: {str(e)}")

    if 'response_content' in st.session_state:
        response_content = st.session_state['response_content']
        
        # Display the response from the model
        # st.markdown("### AI Response:")
        st.markdown('<h3 class="ai-response-heading">AI Response:</h3>', unsafe_allow_html=True)
        st.markdown(response_content)
        
        # Create and download Word document
        word_data = create_word_document(response_content)
        st.download_button(
            label="Download Response as Word",
            data=word_data,
            file_name="response.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Create and download HTML content
        html_content = create_html_content(response_content)
        st.download_button(
            label="Download Response as HTML",
            data=html_content,
            file_name="response.html",
            mime="text/html"
        )

if __name__ == "__main__":
    main()